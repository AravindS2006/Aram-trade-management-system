import docx
import sys

doc_path = r"c:\Users\aravi\Documents\GitHub\Aram-trade-management-system\strategies\weapon_candle\weapon_candle_copilot_instructions.docx"

try:
    doc = docx.Document(doc_path)
    text = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text.append(" | ".join(row_text))
                
    with open(r"c:\Users\aravi\Documents\GitHub\Aram-trade-management-system\extracted_docx.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(text))
    print("Extraction successful.")
except Exception as e:
    print(f"Error: {e}")
